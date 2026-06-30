"""生成 Memo 项目完整说明文档 (DOCX)。"""
import sys; import os; sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from memo.core.engine import engine

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)

# ── 封面 ──
for _ in range(6): doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Memo（麦默）')
run.font.size = Pt(36); run.bold = True; run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('活的、会进化的 Agent 记忆系统')
run.font.size = Pt(16); run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run('v0.1.0  |  MIT License  |  2026-06-27').font.size = Pt(12)

doc.add_page_break()

# ── 目录 ──
doc.add_heading('目录', level=1)
toc_items = [
    '1. 项目概述', '2. 系统要求', '3. 核心架构', '4. 技术原理',
    '5. 功能清单', '6. 安装与部署', '7. 使用指南', '8. MCP 工具参考',
    '9. 看板使用', '10. 多 Agent 集成', '11. 性能参考', '12. 常见问题',
    '附录 A：数据库表结构', '附录 B：关键参数表',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── 1. 项目概述 ──
doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph(
    'Memo（麦默）是一个开源的 AI Agent 记忆层。它为 Claude、Cursor、HanaAgent、WorkBuddy 等'
    '任何支持 MCP 协议的 AI Agent 提供持久化的、会自动关联和进化的记忆能力。'
)
doc.add_paragraph(
    '与传统"把对话存进向量数据库"的做法不同，Memo 构建了一张活的记忆网络：特征词之间通过赫布学习'
    '自动建立带权重的有向边，检索时沿边扩散激活，发现散落在不同会话中的间接关联。同时用 Bjork 双强度'
    '遗忘模型保证记忆不会无限膨胀。'
)

p = doc.add_paragraph()
run = p.add_run('设计哲学：')
run.bold = True
doc.add_paragraph('记忆应该是活的（权重会进化），不是静态的索引。', style='List Bullet')
doc.add_paragraph('记忆应该会遗忘（不重要的自然衰减），不是什么都永存。', style='List Bullet')
doc.add_paragraph('记忆应该能关联（跨会话自动串联），不是孤立的片段。', style='List Bullet')

# ── 2. 系统要求 ──
doc.add_heading('2. 系统要求', level=1)

doc.add_heading('2.1 硬件', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['项目', '最低配置', '推荐配置']):
    table.rows[0].cells[i].text = h
for r, (item, min_cfg, rec_cfg) in enumerate([
    ('CPU', '4 核 (i5-4590S+)', '8 核+ (12代 i5+)'),
    ('内存', '8 GB RAM', '16-32 GB RAM'),
    ('存储', '10 GB SSD 可用', '50+ GB SSD'),
]):
    table.rows[r+1].cells[0].text = item
    table.rows[r+1].cells[1].text = min_cfg
    table.rows[r+1].cells[2].text = rec_cfg

doc.add_heading('2.2 软件', level=2)
doc.add_paragraph('Python 3.11 或更高版本', style='List Bullet')
doc.add_paragraph('Windows / macOS / Linux', style='List Bullet')
doc.add_paragraph('首次运行需下载 BGE-small-zh-v1.5 嵌入模型（约 120 MB），之后离线可用', style='List Bullet')
doc.add_paragraph('可选：OpenAI 兼容 API Key（DeepSeek / OpenAI），用于提升特征词提取质量', style='List Bullet')

doc.add_heading('2.3 网络', level=2)
doc.add_paragraph('核心功能零网络依赖（SQLite 本地存储 + BGE 本地嵌入）。可选 LLM 提取需访问 API。')

# ── 3. 核心架构 ──
doc.add_heading('3. 核心架构', level=1)

doc.add_heading('3.1 三层记忆', level=2)
table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(['层级', '类比', '存储', '延迟']):
    table.rows[0].cells[i].text = h
for r, (l, a, s, d) in enumerate([
    ('L0 热记忆体', 'L1 缓存', '高频特征词 + 近期摘要 (内存)', '<10ms'),
    ('L1 温记忆体', '大脑皮层', '特征关系图谱 + 全文索引 (SQLite)', '100-500ms'),
    ('L2 冷存储', '档案库', '完整原文 + 历史事实 (SQLite)', '1-5s'),
]):
    table.rows[r+1].cells[0].text = l; table.rows[r+1].cells[1].text = a
    table.rows[r+1].cells[2].text = s; table.rows[r+1].cells[3].text = d

doc.add_heading('3.2 六维存储模型', level=2)
doc.add_paragraph('会话 (Session)：标题 + 时间范围 + 状态', style='List Bullet')
doc.add_paragraph('特征词 (FeatureTag)：名称 + 双强度权重 + 激活次数 + 向量嵌入', style='List Bullet')
doc.add_paragraph('特征关系 (FeatureRelation)：源→目标 + 赫布边权重 + 共激次数 + 关系类型', style='List Bullet')
doc.add_paragraph('记忆单元 (MemoryUnit)：标题 + 摘要/二级摘要 + 原文 + 双时序 + 类型 + 置信度', style='List Bullet')
doc.add_paragraph('标签提及 (TagMention)：特征词→记忆单元关联 + 相关度评分', style='List Bullet')
doc.add_paragraph('全局快照 (GlobalSnapshot)：定期生成 Agent 对用户的全局认知画像', style='List Bullet')

doc.add_heading('3.3 三通道检索', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['通道', '技术', '速度']):
    table.rows[0].cells[i].text = h
for r, (c, t, s) in enumerate([
    ('① 向量语义', 'BGE-small cosine 相似度', '~50ms'),
    ('② 全文关键词', 'SQLite FTS5 BM25', '~30ms'),
    ('③ 图扩散激活', '赫布边 BFS 游走（独家）', '~100ms'),
]):
    table.rows[r+1].cells[0].text = c; table.rows[r+1].cells[1].text = t; table.rows[r+1].cells[2].text = s
doc.add_paragraph('三通道结果用 RRF（Reciprocal Rank Fusion）融合排序，返回 Top-K。')

# ── 4. 技术原理 ──
doc.add_heading('4. 技术原理', level=1)

doc.add_heading('4.1 赫布学习', level=2)
doc.add_paragraph(
    '灵感来自神经科学：两个特征词在同一上下文中共同激活时，它们之间的边权重自动增长。'
    '公式：Δw = 0.05 × (1 − w) × 语义相似度 × min(共激活次数/10, 2.0)。'
    '这意味着：越是经常一起出现的特征词，连接越强；连接越强，检索时越容易被扩散激活触发。'
)

doc.add_heading('4.2 扩散激活', level=2)
doc.add_paragraph(
    '从查询中提取种子特征词 → 沿赫布边 BFS 游走（最多 3 跳，每跳衰减 50%）→ '
    '激活的邻居节点继续传播 → 最终着陆到关联的记忆单元。'
    '因果边（CAUSAL）传播系数 2.0×，时序边（TEMPORAL）1.2×，保证了关联传播的语义合理性。'
)

doc.add_heading('4.3 Bjork 双强度遗忘', level=2)
doc.add_paragraph('存储强度（只增不减）：表征"这个信息有多重要"。每次激活 +0.01。', style='List Bullet')
doc.add_paragraph('检索强度（不用时衰减）：表征"现在还能多快想起它"。每天衰减 2%。', style='List Bullet')
doc.add_paragraph('三维度休眠：同时满足有效权重<0.05、休眠>90天、近30天无访问，标记休眠但不删除。', style='List Bullet')

doc.add_heading('4.4 双时序事实追踪', level=2)
doc.add_paragraph(
    '每条记忆携带 valid_from 和 valid_until 字段。新事实与旧事实冲突时，'
    '旧事实自动标记 is_superseded=true，保留完整审计链。支持"以3月1日为基准"的时点查询。'
)

# ── 5. 功能清单 ──
doc.add_heading('5. 功能清单', level=1)
features = [
    ('✅ 记忆写入', '自动特征词提取（LLM/jieba 双模式）、摘要生成、关系建立'),
    ('✅ 三通道检索', '向量语义 + BM25 全文 + 图扩散激活, RRF 融合排序'),
    ('✅ 赫布学习', '特征词共激活时边权重自动增长'),
    ('✅ 双强度遗忘', 'Bjork 模型：重要的留、不用的自然衰减'),
    ('✅ 双时序处理', '事实随时间变化自动标记旧版本'),
    ('✅ 全局快照', '定期生成 Agent 对用户的认知画像'),
    ('✅ MCP Server', '8 个标准 MCP 工具，任何 Agent 可接入'),
    ('✅ Web 看板', '可视化统计、特征词云、记忆列表、详情查看'),
    ('✅ 自动同步', '守护进程监控会话文件，实时同步到 Memo'),
    ('✅ 零 API 依赖', '无 API Key 也能完整运行 (jieba 降级 + 本地嵌入)'),
]
for name, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(name + '：')
    run.bold = True
    p.add_run(desc)

# ── 6. 安装与部署 ──
doc.add_heading('6. 安装与部署', level=1)

doc.add_heading('6.1 获取项目', level=2)
doc.add_paragraph('git clone https://github.com/adoublegirl-dev/memo.git', style='No Spacing')
doc.add_paragraph('cd memo/memo-project', style='No Spacing')

doc.add_heading('6.2 安装依赖', level=2)
doc.add_paragraph('pip install -r requirements.txt', style='No Spacing')

doc.add_heading('6.3 配置环境', level=2)
doc.add_paragraph('cp .env.example .env（复制配置模板）', style='No Spacing')
doc.add_paragraph('编辑 .env，填入 API Key（可选）：', style='No Spacing')
doc.add_paragraph('OPENAI_API_KEY=sk-你的key', style='No Spacing')
doc.add_paragraph('OPENAI_BASE_URL=https://api.deepseek.com/v1', style='No Spacing')
doc.add_paragraph('MEMO_EXTRACTION_MODEL=deepseek-v4-flash', style='No Spacing')

doc.add_heading('6.4 验证安装', level=2)
doc.add_paragraph('python scripts/quick_check.py', style='No Spacing')
doc.add_paragraph('看到 "Phase 0 基础设施验证通过！" 即成功。', style='No Spacing')

doc.add_heading('6.5 接入 Agent', level=2)
doc.add_paragraph(
    '在任何 MCP 兼容 Agent 的配置文件中添加 Memo 连接器。以 Claude Desktop 为例：'
)
doc.add_paragraph(
    '编辑 %APPDATA%\\Claude\\claude_desktop_config.json：'
)
doc.add_paragraph(
    '{"mcpServers": {"memo": {"command": "python", "args": ["E:/memo/scripts/run_mcp.py"]}}}',
    style='No Spacing'
)

doc.add_heading('6.6 启动后台服务', level=2)
doc.add_paragraph('双击 E:\\memo\\start_all.bat 一键启动看板 + 自动同步。', style='List Bullet')
doc.add_paragraph('双击 E:\\memo\\stop_all.bat 停止所有后台服务。', style='List Bullet')

# ── 7. 使用指南 ──
doc.add_heading('7. 使用指南', level=1)

doc.add_heading('7.1 对话中写入记忆', level=2)
doc.add_paragraph('方式一（手动）：对 Agent 说「帮我记住 XXX」', style='List Bullet')
doc.add_paragraph('方式二（自动）：启动 memo_watcher.py，每轮对话自动同步', style='List Bullet')

doc.add_heading('7.2 对话中检索记忆', level=2)
doc.add_paragraph('对 Agent 说「之前 XXX 是什么来着？」或「还记得 XXX 吗？」', style='List Bullet')
doc.add_paragraph('Agent 会自动调用 memo_recall 搜索相关记忆并注入上下文。', style='List Bullet')

doc.add_heading('7.3 日常维护', level=2)
doc.add_paragraph('对 Agent 说「给 Memo 做一次维护」触发生命周期（遗忘+固化+快照）。', style='List Bullet')
doc.add_paragraph('也可直接调 memo_maintain 工具。', style='List Bullet')

# ── 8. MCP 工具参考 ──
doc.add_heading('8. MCP 工具参考', level=1)
tools = [
    ('memo_remember', '写入记忆', '提供 conversation 自动提取特征词/摘要/关系'),
    ('memo_recall', '三通道检索', 'query 查询文本, top_k 返回数量'),
    ('memo_start_session', '开始会话', 'title 会话标题'),
    ('memo_end_session', '结束会话', '结束当前活跃会话'),
    ('memo_stats', '统计信息', '返回会话数/记忆数/特征词数/关系数/高频词'),
    ('memo_hot_tags', '高频特征词', 'limit 返回数量, 了解 Agent 当前活跃知识域'),
    ('memo_maintain', '生命周期维护', '遗忘衰减 + 固化 + 快照'),
    ('memo_snapshot', '全局快照', 'Agent 对用户的最新认知画像'),
]
table = doc.add_table(rows=len(tools)+1, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['工具名', '功能', '参数']):
    table.rows[0].cells[i].text = h
for r, (name, func, params) in enumerate(tools):
    table.rows[r+1].cells[0].text = name; table.rows[r+1].cells[1].text = func
    table.rows[r+1].cells[2].text = params

# ── 9. 看板使用 ──
doc.add_heading('9. 看板使用', level=1)
doc.add_paragraph('启动后访问 http://localhost:9120，功能包括：')
doc.add_paragraph('统计卡片：实时显示会话数、记忆数、特征词数、关系数、向量索引大小', style='List Bullet')
doc.add_paragraph('特征词云：按权重着色/调大小，点击任意标签即搜索', style='List Bullet')
doc.add_paragraph('记忆列表：支持关键词搜索、类型筛选（FACT/DECISION/PREFERENCE/EVENT/REASONING）', style='List Bullet')
doc.add_paragraph('详情弹窗：点击记忆卡片查看完整原文和二级摘要', style='List Bullet')

# ── 10. 多 Agent 集成 ──
doc.add_heading('10. 多 Agent 集成', level=1)
doc.add_paragraph(
    'Memo 通过 MCP 协议对接 Agent，只要是 MCP 兼容的 Agent 都能接入。'
    '目前已在 HanaAgent 上完整验证：MCP Server 直连 + 守护进程自动同步 + 看板可视化。'
)
doc.add_paragraph(
    '对于 Claude Desktop / Cursor / Claude Code / WorkBuddy 等，只需在对应的 MCP 配置文件中'
    '添加 Memo 连接器即可。所有 Agent 可共享同一个数据库文件，特征词图谱跨 Agent 联动。'
)

# ── 11. 性能参考 ──
doc.add_heading('11. 性能参考', level=1)
stats = engine.stats()
doc.add_paragraph(f'当前数据规模：{stats["sessions"]} 会话, {stats["memories"]} 记忆, '
                  f'{stats["feature_tags"]} 特征词, {stats["relations"]} 关系')
doc.add_paragraph('SQLite 单文件存储，1 万次对话约 50-100 MB。', style='List Bullet')
doc.add_paragraph('向量索引驻留内存，10 万条以下无需 GPU。', style='List Bullet')
doc.add_paragraph('图遍历为 BFS 邻接表，1 万特征词 + 10 万关系时单次检索 < 200ms。', style='List Bullet')
doc.add_paragraph('数据量突破 SQLite 边界时可平滑迁移到 Neo4j（架构预留接口）。', style='List Bullet')

# ── 12. 常见问题 ──
doc.add_heading('12. 常见问题', level=1)
faq = [
    ('Q: 没有 API Key 能用吗？', 'A: 能。Memo 零 API 依赖运行。API Key 只是让特征词提取从 jieba 关键词升级为 LLM 语义提取，质量更高。'),
    ('Q: 数据安全吗？', 'A: 全部数据存储在本地的单个 SQLite 文件中（E:\\memo\\memo\\data\\memo.db）。无云服务、无遥测、无第三方传输。LLM 提取时只传当前对话片段（不传全量历史）。'),
    ('Q: 数据库会无限膨胀吗？', 'A: 不会。Bjork 双强度遗忘自动衰减不活跃的记忆。1000 次对话约 10 MB。'),
    ('Q: 能多个 Agent 共用一个记忆库吗？', 'A: 可以。所有 Agent 的 MCP 配置指向同一个数据库路径即可。特征词图谱跨 Agent 共享。'),
    ('Q: 支持英文吗？', 'A: BGE-small-zh-v1.5 中英文都支持。英文场景可切换到 BGE-base-en-v1.5。'),
    ('Q: 怎么备份？', 'A: 复制 E:\\memo\\memo\\data\\memo.db 文件即可。SQLite 单文件，随时复制。'),
]
for q, a in faq:
    p = doc.add_paragraph()
    run = p.add_run(q + ' ')
    run.bold = True
    p.add_run(a)

# ── 附录 ──
doc.add_page_break()
doc.add_heading('附录 A：数据库表结构', level=1)
doc.add_paragraph('sessions — 会话（id, agent_id, title, status, created_at, memory_count）')
doc.add_paragraph('feature_tags — 特征词（id, name, category, storage_strength, retrieval_strength, total_activations, embedding）')
doc.add_paragraph('feature_relations — 特征关系（source_tag_id, target_tag_id, relation_type, hebbian_weight, co_activation_count）')
doc.add_paragraph('memory_units — 记忆单元（session_id, title, summary, summary_detail, raw_text, valid_from, valid_until, is_superseded, memory_type, embedding）')
doc.add_paragraph('tag_mentions — 标签提及（tag_id, memory_unit_id, mention_type, relevance_score）')
doc.add_paragraph('global_snapshots — 全局快照（agent_profile, top_domains, hot_tags, active_projects）')
doc.add_paragraph('memory_fts — FTS5 全文搜索索引')

doc.add_heading('附录 B：关键参数表', level=1)
params = [
    ('hebbian_learning_rate', '0.05', '赫布学习率'),
    ('spreading_decay_rate', '0.5', '扩散激活每跳衰减率'),
    ('spreading_max_hops', '3', '最大扩散跳数'),
    ('retrieval_strength_decay', '0.02/天', '检索强度日衰减率'),
    ('storage_strength_increment', '0.01/次', '每次激活存储强度增量'),
    ('dormant_threshold_days', '90', '休眠天数阈值'),
    ('consolidation_trigger_count', '10', '固化触发数'),
    ('snapshot_trigger_count', '50', '快照触发数'),
    ('embedding_dim', '512', 'BGE-small 嵌入维度'),
    ('top_k_retrieval', '5', '默认检索返回数'),
]
table = doc.add_table(rows=len(params)+1, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['参数', '默认值', '说明']):
    table.rows[0].cells[i].text = h
for r, (name, val, desc) in enumerate(params):
    table.rows[r+1].cells[0].text = name; table.rows[r+1].cells[1].text = val
    table.rows[r+1].cells[2].text = desc

# ── 保存 ──
output_path = "E:/memo/docs/Memo-项目说明文档.docx"
doc.save(output_path)
print(f"✅ 文档已生成: {output_path}")
